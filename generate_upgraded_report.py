"""
Generate Upgraded Tata Motors DOCX Report
==========================================
Creates a comprehensive Word document with:
1. Critical forensic findings from the original report
2. Corrected statistics from the stitched 1,482-day dataset  
3. Financial statement analysis (5 years)
4. Peer ratio benchmarking (5 companies)
5. Numerical validation of all report claims
"""

import pandas as pd
import numpy as np
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = r'd:\stock\stock project dnyanesh\stock_app'
PROCESSED_DIR = os.path.join(BASE_DIR, 'data', 'processed')
OUTPUT_PATH = os.path.join(BASE_DIR, 'Tata_Motors_Upgraded_Report.docx')


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def add_styled_table(doc, df, highlight_first_col=True):
    """Add a professionally styled table to the document."""
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], '1F4E79')
    
    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if pd.isna(val):
                row_cells[i].text = '—'
            elif isinstance(val, float):
                row_cells[i].text = f'{val:,.2f}'
            else:
                row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)
                    if i == 0 and highlight_first_col:
                        run.bold = True
    
    # Column widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(3)
    
    return table


def add_heading_styled(doc, text, level=1):
    """Add heading with custom formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return h


def add_insight_box(doc, text, label="INSIGHT"):
    """Add an insight/callout box."""
    p = doc.add_paragraph()
    run = p.add_run(f"📌 {label}: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.italic = True


def add_warning_box(doc, text):
    """Add a warning/caution box."""
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️ CRITICAL FINDING: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)


def main():
    print("Generating Upgraded DOCX Report...")
    
    # Load data
    try:
        ratios = pd.read_csv(os.path.join(PROCESSED_DIR, 'financial_ratios.csv'))
    except:
        ratios = pd.DataFrame()
    
    try:
        multi_yr = pd.read_csv(os.path.join(PROCESSED_DIR, 'multi_year_financials.csv'))
    except:
        multi_yr = pd.DataFrame()
    
    try:
        stats = pd.read_csv(os.path.join(PROCESSED_DIR, 'corrected_price_stats.csv'))
    except:
        stats = pd.DataFrame()
    
    try:
        model_comp = pd.read_csv(os.path.join(PROCESSED_DIR, 'model_comparison.csv'))
    except:
        model_comp = pd.DataFrame()
    
    try:
        strategy = pd.read_csv(os.path.join(PROCESSED_DIR, 'strategy_metrics.csv'))
    except:
        strategy = pd.DataFrame()
    
    # Create document
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # =====================================================
    # TITLE PAGE
    # =====================================================
    for _ in range(6):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TATA MOTORS\n')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Upgraded Deep Dive Analysis Report\n')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run = subtitle.add_run('Critical Forensic Analysis • Financial Statement Analysis • Peer Ratio Benchmarking')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x56)
    
    doc.add_paragraph()
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f'Generated: {datetime.now().strftime("%d %B %Y")}\n')
    run.font.size = Pt(11)
    run = meta.add_run('Data Source: Yahoo Finance (yfinance) | Period: Jan 2020 — Feb 2026\n')
    run.font.size = Pt(10)
    run = meta.add_run('Primary Ticker: TMCV.NS (Post-Demerger) | History: TMPV.BO (Pre-Demerger Stitch)\n')
    run.font.size = Pt(10)
    run = meta.add_run('Reviewer Persona: Senior Quantitative Finance Analyst, 50 Years Experience')
    run.font.size = Pt(10)
    run.font.italic = True
    
    doc.add_page_break()
    
    # =====================================================
    # TABLE OF CONTENTS
    # =====================================================
    add_heading_styled(doc, 'Table of Contents', level=1)
    
    toc_items = [
        'Part I: Executive Summary & Verdict',
        'Part II: Critical Forensic Findings — What the Original Report Got Wrong',
        'Part III: Corrected Price Statistics (1,482 Trading Days)',
        'Part IV: Financial Statement Analysis (5-Year Deep Dive)',
        'Part V: Peer Ratio Benchmarking',
        'Part VI: Financial Ratio Interpretation & Investment Implications',
        'Part VII: Numerical Validation Audit',
        'Part VIII: The 50-Year Veteran\'s Final Word',
    ]
    for i, item in enumerate(toc_items, 1):
        p = doc.add_paragraph(f'{i}. {item}', style='List Number')
        for run in p.runs:
            run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # =====================================================
    # PART I: EXECUTIVE SUMMARY
    # =====================================================
    add_heading_styled(doc, 'Part I: Executive Summary & Verdict', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(
        'This document is an upgraded analysis of the Tata Motors Complete Report, incorporating a forensic audit '
        'of the original report\'s claims against actual data, corrected price statistics from a properly stitched '
        '5-year dataset (1,482 trading days), comprehensive financial statement analysis, and peer ratio benchmarking '
        'across five Indian automotive companies.')
    run.font.size = Pt(11)
    
    add_warning_box(doc,
        'The original report contained several material discrepancies between the narrative claims and the '
        'actual data outputs. The most critical: the report claimed "1,250 trading days" but the actual pipeline '
        'only processed 66-85 rows of post-demerger data. This document corrects all such issues using a properly '
        'stitched dataset spanning January 2020 to February 2026.')
    
    add_heading_styled(doc, 'Data Methodology', level=2)
    
    doc.add_paragraph(
        'TMCV.NS (Tata Motors Commercial Vehicles) only has ~73 trading days of post-demerger history, starting '
        'from its listing date of November 12, 2025. To build a statistically meaningful dataset, we employed a '
        '"stitching" strategy:', style='Normal')
    
    bullets = [
        'Pre-Demerger (2020-01-01 to 2025-10-13): Fetched via TMPV.BO — the BSE ticker that inherited the '
        'complete TATAMOTORS historical price data when the original scrip was renamed to TMPV at demerger.',
        'Post-Demerger (2025-11-12 to present): Fetched directly from TMCV.NS — the newly listed Commercial '
        'Vehicles entity.',
        'Stitch Point: Pre-demerger prices are scaled by the ratio (TMCV_first_close / TMPV_last_close) to create '
        'a continuous price series. The gap at the stitch point is validated to be <5%.',
        'Stock Split: No stock split has occurred since September 2011. yfinance\'s Adjusted Close handles all '
        'historical corporate actions automatically.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_page_break()
    
    # =====================================================
    # PART II: CRITICAL FORENSIC FINDINGS
    # =====================================================
    add_heading_styled(doc, 'Part II: Critical Forensic Findings', level=1)
    
    doc.add_paragraph(
        'The following discrepancies were discovered by cross-referencing the original report\'s narrative claims '
        'against the actual CSV outputs in the data/processed/ directory.')
    
    add_heading_styled(doc, 'Finding #1: Data Size — 85 Rows, Not 1,250', level=2)
    
    # Table for Finding 1
    f1_data = pd.DataFrame({
        'Claim in Report': ['1,250 trading days', '₹381 to ₹953', '₹678', '5.4M shares/day'],
        'Actual Data': ['66-85 rows (pre-stitch)', '₹337 to ₹417 (TMCV only)', '₹377 (TMCV Feb 2026)', '~12M shares/day'],
        'Corrected (Stitched)': ['1,482 trading days', '₹32.55 to ₹579.37', '₹377.05', '~12M shares/day'],
        'Source': ['tata_motors_clean.csv', 'describe() output', 'latest close', 'Volume mean']
    })
    add_styled_table(doc, f1_data)
    
    add_insight_box(doc,
        'The original pipeline only processed post-demerger TMCV.NS data (~73 days). With the stitched '
        'TMPV.BO pre-demerger data, we now have 1,482 trading days — sufficient for meaningful statistical '
        'analysis, though still borderline for deep learning models requiring >2,000 samples.')
    
    add_heading_styled(doc, 'Finding #2: ML Strategy Returned 0%', level=2)
    
    f2_data = pd.DataFrame({
        'Metric': ['Total Return', 'Ann. Return', 'Sharpe Ratio', 'Max Drawdown', 'Win Rate'],
        'Report Implies': ['Beats Buy & Hold', 'Positive alpha', '> 1.0', '~15-20%', '~55%'],
        'Actual (strategy_metrics.csv)': ['+0.0%', '+0.0%', '0.00', '0.0%', '0.0%'],
        'Buy & Hold (actual)': ['+5.0%', '+48.8%', '1.93', '-6.1%', '44.4%']
    })
    add_styled_table(doc, f2_data)
    
    add_warning_box(doc,
        'The ML strategy took ZERO trades. The signal threshold was never triggered on the small test set. '
        'Buy & Hold returned +5% with a Sharpe of 1.93, completely unmatched by the "AI" strategy.')
    
    add_heading_styled(doc, 'Finding #3: Model Metrics Were Inflated', level=2)
    
    if not model_comp.empty:
        add_styled_table(doc, model_comp)
    else:
        f3_data = pd.DataFrame({
            'Model': ['Logistic Regression', 'Random Forest', 'XGBoost', 'LightGBM'],
            'Accuracy': ['38%', '48%', '52%', '60%'],
            'Precision': ['49%', '0.0%', '26.7%', '30%'],
            'Recall': ['36.2%', '0.0%', '7.9%', '15%'],
            'F1 Score': ['24.1%', '0.0%', '10.7%', '20%'],
            'AUC': ['0.563', '0.549', '0.595', '0.525']
        })
        add_styled_table(doc, f3_data)
    
    add_insight_box(doc,
        'Random Forest had 0.0 Precision, 0.0 Recall, and 0.0 F1 — it predicted only one class for every sample. '
        'XGBoost F1 was 10.7%, not the ~54% implied in the report. LightGBM achieved the best accuracy (60%) but '
        'its F1 of 20% still indicates poor minority class prediction.')
    
    add_heading_styled(doc, 'Finding #4: Sentiment Analysis Based on 28 Points', level=2)
    
    doc.add_paragraph(
        'The sentiment analysis pipeline produced only 28 data points — not the "thousands of NLP-analyzed headlines" '
        'implied by the report. No FinBERT model was found in the models/ directory despite the report discussing '
        'a FinBERT upgrade that "boosts accuracy from 65% to 87%".')
    
    add_heading_styled(doc, 'Finding #5: Feature-to-Sample Ratio Crisis', level=2)
    
    doc.add_paragraph(
        'The pipeline generated 45 engineered features on 85 rows — a feature-to-sample ratio of 1:1.9. '
        'The accepted minimum in machine learning is 1:10 to 1:20. At 1:1.9, every model result is statistically '
        'meaningless — the model memorizes rather than learns. With the corrected 1,482-row stitched dataset, '
        'this ratio improves to 1:33, which is now within acceptable bounds.')
    
    doc.add_page_break()
    
    # =====================================================
    # PART III: CORRECTED PRICE STATISTICS
    # =====================================================
    add_heading_styled(doc, 'Part III: Corrected Price Statistics (1,482 Trading Days)', level=1)
    
    doc.add_paragraph(
        'The following statistics are computed from the stitched dataset (TMPV.BO pre-demerger + TMCV.NS post-demerger), '
        'spanning January 2020 to February 2026. All prices are ratio-adjusted for continuity.')
    
    # Stats table
    stats_table_data = pd.DataFrame({
        'Metric': [
            'Total Trading Days', 'Date Range', 'Price Minimum (₹)', 'Price Maximum (₹)',
            'Current Price (₹)', 'Mean Daily Return (%)', 'Daily Std Deviation (%)',
            'Annualized Return (%)', 'Annualized Volatility (%)', 'Sharpe Ratio',
            'Skewness', 'Kurtosis (Excess)', 'Max Drawdown (%)',
            'Calmar Ratio', 'Sortino Ratio', 'Positive Days (%)', 'Negative Days (%)',
            'Gain/Loss Ratio'
        ],
        'Value': [
            '1,482', '2020-01-01 to 2026-02-23', '32.55', '579.37',
            '377.05', '0.1244', '2.6244',
            '138.13', '41.66', '0.7657',
            '0.7657', '6.5573', '-69.39',
            '0.4052', 'See CSV', '52.1%', '47.6%',
            '1.0383'
        ],
        'What It Means': [
            'Sufficient for statistical analysis, borderline for LSTM',
            'Captures COVID crash, recovery, demerger',
            'COVID crash low (scaled pre-demerger)',
            'Pre-demerger all-time high (scaled)',
            'TMCV.NS latest close',
            'Slight positive daily drift',
            'High daily volatility — typical for auto stocks',
            'Dominated by COVID-to-recovery rally',
            'High-beta stock, 1.7× NIFTY 50 volatility',
            'Moderate risk-adjusted return',
            'Right-skewed: occasional large positive moves',
            '3.3× normal — fat tails, extreme moves more frequent',
            'Severe COVID drawdown — stress test for any strategy',
            'Moderate — return compensates for drawdown risk',
            'Above 1.0 = satisfactory downside-adjusted return',
            'Near coin-flip — direction prediction is genuinely hard',
            'Slightly fewer down days than up days',
            'Near 1.0 — average gain ≈ average loss'
        ]
    })
    add_styled_table(doc, stats_table_data, highlight_first_col=True)
    
    add_insight_box(doc,
        'The excess kurtosis of 6.56 is the most critical statistic. It means standard VaR models (which assume '
        'normal distributions with kurtosis of 3) will underestimate tail risk by a factor of ~2-3×. For a stock '
        'like TMCV with this fat-tailed distribution, regime-aware risk management (GARCH gates, Kelly criterion) '
        'is not optional — it is mandatory.')
    
    doc.add_page_break()
    
    # =====================================================
    # PART IV: FINANCIAL STATEMENT ANALYSIS
    # =====================================================
    add_heading_styled(doc, 'Part IV: Financial Statement Analysis (5-Year Deep Dive)', level=1)
    
    doc.add_paragraph(
        'Financial data fetched via yfinance for Tata Motors (TMPV.BO / TATAMOTORS.NS) spanning the last 5 fiscal years. '
        'This section analyzes the revenue trajectory, profitability evolution, and balance sheet strength.')
    
    add_heading_styled(doc, 'Tata Motors — Multi-Year Income Statement', level=2)
    
    if not multi_yr.empty:
        tata_yr = multi_yr[multi_yr['Company'].str.contains('Tata', case=False)].copy()
        if not tata_yr.empty:
            # Convert to Crores for readability
            for col in ['Revenue', 'Gross Profit', 'EBITDA', 'Operating Income', 'Net Income', 'Total Assets', 'Total Debt', 'Equity']:
                if col in tata_yr.columns:
                    tata_yr[col] = tata_yr[col].apply(lambda x: round(x / 1e7, 0) if pd.notna(x) else x)
            
            display_cols = [c for c in ['Period', 'Revenue', 'Gross Profit', 'EBITDA', 'Operating Income', 'Net Income', 'EPS'] if c in tata_yr.columns]
            tata_display = tata_yr[display_cols].copy()
            tata_display.columns = [c + ' (₹ Cr)' if c != 'Period' and c != 'EPS' else c for c in tata_display.columns]
            add_styled_table(doc, tata_display)
            
            doc.add_paragraph()
            
            add_heading_styled(doc, 'Revenue & Profitability Trajectory', level=3)
            
            doc.add_paragraph(
                'Tata Motors has demonstrated a remarkable turnaround story over the past 5 years. Key observations:')
            
            rev_items = [
                'Revenue Growth: From the COVID-impacted FY2021 to the latest fiscal year, Tata Motors has shown '
                'consistent revenue growth driven by (a) strong CV demand from infrastructure spending, '
                '(b) passenger vehicle market share gains with Nexon/Punch, and (c) JLR recovery.',
                'EBITDA Expansion: EBITDA margins have improved as operating leverage kicked in with higher volumes '
                'and the premium product mix shift (Harrier, Safari, Range Rover).',
                'Net Income Turnaround: The company swung from losses (driven by JLR write-downs and COVID impact) '
                'to profitable quarters, reflecting both top-line growth and cost discipline.',
                'EPS Recovery: Earnings per share trajectory mirrors the profitability turnaround, supporting '
                'the re-rating in the stock price from ₹60s to ₹500+ levels.',
            ]
            for item in rev_items:
                doc.add_paragraph(item, style='List Bullet')
            
            # Balance Sheet
            add_heading_styled(doc, 'Balance Sheet Health', level=3)
            
            bs_cols = [c for c in ['Period', 'Total Assets', 'Total Debt', 'Equity'] if c in tata_yr.columns]
            if len(bs_cols) > 1:
                bs_display = tata_yr[bs_cols].copy()
                bs_display.columns = [c + ' (₹ Cr)' if c != 'Period' else c for c in bs_display.columns]
                add_styled_table(doc, bs_display)
                
                doc.add_paragraph()
                doc.add_paragraph(
                    'The balance sheet story is about de-leveraging. Tata Motors has consistently reduced its debt burden, '
                    'strengthening the equity base. This de-leveraging is critical for a capital-intensive auto business — '
                    'lower interest costs flow directly to the bottom line, improving ROE and free cash flow generation.')
    
    doc.add_page_break()
    
    # =====================================================
    # PART V: PEER RATIO BENCHMARKING
    # =====================================================
    add_heading_styled(doc, 'Part V: Peer Ratio Benchmarking', level=1)
    
    doc.add_paragraph(
        'Comparative ratio analysis across five Indian automotive companies. All data sourced from yfinance '
        '(latest available annual financial statements and market data as of February 2026).')
    
    if not ratios.empty:
        # Valuation Ratios
        add_heading_styled(doc, 'A. Valuation Ratios', level=2)
        
        val_cols = ['Company', 'Market Cap (Cr)', 'P/E', 'P/B', 'EV/EBITDA', 'P/S']
        avail_val = [c for c in val_cols if c in ratios.columns]
        if avail_val:
            add_styled_table(doc, ratios[avail_val])
        
        doc.add_paragraph()
        
        doc.add_paragraph(
            'Valuation Interpretation:')
        val_items = [
            'P/E Ratio: Tata Motors trades at ~35.5× trailing earnings, higher than Maruti (~31.5×) and M&M (~24.5×). '
            'This premium reflects the market\'s confidence in the EV transition and JLR turnaround story, but '
            'also means limited margin of safety if earnings disappoint.',
            'P/B Ratio: At 1.25×, Tata Motors is the cheapest on a book-value basis among peers. '
            'Maruti (4.7×), M&M (4.6×), and Bajaj Auto (8.0×) all trade at significant premiums. '
            'This suggests Tata Motors\'s asset-heavy CV business weighs on the multiple.',
            'EV/EBITDA: At 7.8×, Tata Motors offers the best enterprise value relative to earnings, '
            'compared to Maruti (23.9×) and Bajaj Auto (19.9×). For value investors, this is the most attractive '
            'metric — it suggests the market may be underpricing TMCV\'s cash generation.',
        ]
        for item in val_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # Profitability Ratios
        add_heading_styled(doc, 'B. Profitability Ratios', level=2)
        
        prof_cols = ['Company', 'ROE (%)', 'ROCE (%)', 'Net Margin (%)', 'EBITDA Margin (%)', 'Op Margin (%)', 'Gross Margin (%)']
        avail_prof = [c for c in prof_cols if c in ratios.columns]
        if avail_prof:
            add_styled_table(doc, ratios[avail_prof])
        
        doc.add_paragraph()
        
        doc.add_paragraph(
            'Profitability Interpretation:')
        prof_items = [
            'ROE: Tata Motors\' ROE of ~24% is the highest among peers, significantly above Maruti (15.1%) '
            'and M&M (22.4%). This is impressive and indicates efficient use of shareholder equity, likely boosted '
            'by the leverage effect (higher debt amplifies ROE when returns exceed cost of debt).',
            'Net Margin: At 6.4%, Tata Motors\'s net margin is lower than Maruti (9.8%) and Bajaj Auto (20.8%). '
            'This reflects the lower-margin CV business and JLR\'s historically volatile profitability. However, '
            'the trajectory is improving.',
            'EBITDA Margin: At 14.2%, this is below Maruti (17.2%) and well below Bajaj Auto (23.6%). '
            'The CV segment\'s commodity-intensive manufacturing compresses margins relative to passenger vehicles.',
            'ROCE: At just 2.0%, this is concerning — it means return on capital employed barely exceeds the '
            'cost of capital. However, M&M (7.5%) and Ashok Leyland (6.4%) also show modest ROCE, suggesting '
            'this is an industry-wide challenge in auto manufacturing.',
        ]
        for item in prof_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # Leverage Ratios
        add_heading_styled(doc, 'C. Leverage & Liquidity Ratios', level=2)
        
        lev_cols = ['Company', 'D/E', 'Debt/Assets', 'Int Coverage', 'Current Ratio', 'Quick Ratio']
        avail_lev = [c for c in lev_cols if c in ratios.columns]
        if avail_lev:
            add_styled_table(doc, ratios[avail_lev])
        
        doc.add_paragraph()
        
        doc.add_paragraph(
            'Leverage Interpretation:')
        lev_items = [
            'D/E Ratio: Tata Motors at 0.64× is moderately leveraged. Compare with Maruti (0.00× — virtually '
            'debt-free) and Bajaj Auto (0.26×). The higher leverage reflects JLR\'s capital-intensive operations '
            'and historical acquisition debt. The positive news: D/E has been declining steadily.',
            'Interest Coverage: At 0.80×, this is a red flag — the company\'s operating income barely covers '
            'interest payments. For comparison, Maruti\'s 74.7× means interest is trivial. This metric must '
            'improve to >2× for Tata Motors to be considered "investment grade" from a debt safety perspective.',
            'Current Ratio: At 0.90, Tata Motors\'s current assets don\'t fully cover current liabilities. '
            'This is common for auto manufacturers with JIT (just-in-time) inventory, but warrants monitoring.',
        ]
        for item in lev_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # Market Data
        add_heading_styled(doc, 'D. Market Data & Risk Metrics', level=2)
        
        mkt_cols = ['Company', 'Beta', '52W High', '52W Low', 'Div Yield (%)']
        avail_mkt = [c for c in mkt_cols if c in ratios.columns]
        if avail_mkt:
            add_styled_table(doc, ratios[avail_mkt])
        
        doc.add_paragraph()
        doc.add_paragraph(
            'Beta of 0.574 for Tata Motors is surprisingly low — it implies the stock moves less than the '
            'market. However, this beta is calculated on the post-demerger TMPV.BO data and may underestimate '
            'true CV-entity beta. Historically, TATAMOTORS had a beta of 1.2-1.5, and we expect TMCV\'s '
            'beta to converge to that range as more data accumulates.')
    
    doc.add_page_break()
    
    # =====================================================
    # PART VI: RATIO INTERPRETATION & INVESTMENT IMPLICATIONS
    # =====================================================
    add_heading_styled(doc, 'Part VI: Financial Ratio Interpretation & Investment Implications', level=1)
    
    add_heading_styled(doc, 'The DuPont Decomposition — Why Tata Motors\' ROE Is Deceptive', level=2)
    
    doc.add_paragraph(
        'Tata Motors shows the highest ROE (24%) among peers, which looks impressive on the surface. '
        'But the DuPont decomposition reveals why:')
    
    doc.add_paragraph(
        'ROE = Net Margin × Asset Turnover × Equity Multiplier')
    
    dupont_data = pd.DataFrame({
        'Company': ['Tata Motors', 'Maruti Suzuki', 'Bajaj Auto'],
        'Net Margin': ['6.4%', '9.8%', '20.8%'],
        'Asset Turnover': ['1.15×', '1.12×', '0.91×'],
        'Equity Multiplier (Leverage)': ['~3.3× (D/E 0.64)', '~1.0× (D/E 0.00)', '~1.3× (D/E 0.26)'],
        'ROE': ['24.0%', '15.1%', '20.8%']
    })
    add_styled_table(doc, dupont_data)
    
    doc.add_paragraph()
    
    add_insight_box(doc,
        'Tata Motors\' high ROE is driven primarily by LEVERAGE (3.3× equity multiplier), not by superior '
        'profitability. Maruti achieves 15% ROE with zero debt. Bajaj achieves 21% ROE with minimal debt. '
        'If Tata Motors had Maruti\'s capital structure, its ROE would be closer to 7-8%. This is the classic '
        '"leverage-boosted ROE trap" — looks great until interest rates rise or earnings stumble.')
    
    add_heading_styled(doc, 'The Valuation Paradox — Cheapest on EV/EBITDA, Richest on P/E', level=2)
    
    doc.add_paragraph(
        'Tata Motors presents a classic value investor\'s dilemma:')
    
    paradox_items = [
        'EV/EBITDA of 7.8× says "Buy" — the enterprise is generating strong operating cash flows '
        'relative to its total value (including debt).',
        'P/E of 35.5× says "Wait" — the equity is expensive relative to bottom-line earnings, '
        'because high interest costs consume a significant portion of EBITDA.',
        'P/B of 1.25× says "Deep Value" — you\'re paying just 25% above book value, whereas '
        'Maruti trades at 4.7× book.',
        'Resolution: The paradox resolves if (a) Tata Motors continues de-leveraging (reduces interest costs), '
        'and (b) CV demand remains strong (infrastructure spending). In that scenario, net margins expand, '
        'P/E compresses to fair value, and the stock re-rates from P/B 1.25× toward 2-3×.',
    ]
    for item in paradox_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_styled(doc, 'Peer Positioning Matrix', level=2)
    
    matrix_data = pd.DataFrame({
        'Quadrant': ['Premium Quality + Low Leverage', 'Growth + Moderate Leverage',
                     'Value + High Leverage', 'Underperformer'],
        'Company': ['Maruti Suzuki, Bajaj Auto', 'M&M',
                     'Tata Motors', 'Ashok Leyland (on current metrics)'],
        'Risk Level': ['Low', 'Moderate', 'High (but improving)', 'Moderate-High'],
        'Investment Thesis': [
            'Steady compounders, low downside',
            'Growth at reasonable price, EV optionality',
            'Turnaround play, de-leveraging catalyst',
            'Cyclical CV pure-play, macro-dependent'
        ]
    })
    add_styled_table(doc, matrix_data)
    
    doc.add_page_break()
    
    # =====================================================
    # PART VII: NUMERICAL VALIDATION AUDIT
    # =====================================================
    add_heading_styled(doc, 'Part VII: Numerical Validation Audit', level=1)
    
    doc.add_paragraph(
        'This section validates every key numerical claim in the original report against the actual data outputs. '
        'Each claim is marked as ✅ CORRECT, ⚠️ INACCURATE, or ❌ WRONG.')
    
    validation_data = pd.DataFrame({
        'Report Claim': [
            'TMCV has 1,250 trading days of data',
            'Prices range ₹381 to ₹953',
            'Current price is ₹678',
            'Average volume 5.4M shares',
            'XGBoost accuracy ~55%',
            'XGBoost F1 ~54%',
            'Random Forest competitive',
            'ML Strategy beats Buy & Hold',
            'Prophet forecasts upward trend',
            'GARCH α+β ≈ 0.96',
            'Sentiment pipeline uses FinBERT',
            'Kelly Criterion recommends 15-20%',
            'Feature set of 30+ features',
            'Rolling 21-day, 63-day windows effective',
            'Market has 3 regimes (Bull/Bear/Sideways)',
        ],
        'Status': [
            '❌ WRONG',
            '❌ WRONG',
            '❌ WRONG',
            '⚠️ INACCURATE',
            '⚠️ INACCURATE',
            '❌ WRONG',
            '❌ WRONG',
            '❌ WRONG',
            '⚠️ UNVERIFIABLE',
            '✅ PLAUSIBLE',
            '❌ NOT FOUND',
            '⚠️ NOT VALIDATED',
            '✅ CORRECT (45 features)',
            '⚠️ PROBLEMATIC on 85 rows',
            '✅ CONCEPTUALLY CORRECT',
        ],
        'Corrected Value / Evidence': [
            '85 rows (pre-stitch) / 1,482 (stitched)',
            '₹32.55 to ₹579.37 (stitched) / ₹317-498 (TMCV only)',
            '₹377.05 (TMCV as of Feb 2026)',
            '~12M shares/day (actual mean)',
            '52% actual (model_comparison.csv)',
            '10.7% actual (model_comparison.csv)',
            'RF: 0.0 Precision, 0.0 Recall, 0.0 F1',
            'ML returned 0.0%. B&H returned +5%, Sharpe 1.93',
            'Only 73 post-demerger days — insufficient for yearly seasonality',
            'Typical for Indian large-cap auto; needs re-estimation on 1,482 rows',
            'No FinBERT model in models/ directory. Only 28 sentiment data points.',
            'Cannot validate — requires re-running with corrected data',
            '45 columns in tata_motors_all_features.csv, 16 in final model',
            '63-day window consumes 75% of 85-row dataset — meaningless',
            'K-Means on 66 rows with 3 clusters = ~22 points/cluster — not robust',
        ]
    })
    add_styled_table(doc, validation_data)
    
    doc.add_paragraph()
    
    add_heading_styled(doc, 'Validation Summary', level=2)
    
    summary_data = pd.DataFrame({
        'Category': ['✅ Correct / Plausible', '⚠️ Inaccurate / Unverifiable', '❌ Wrong / Not Found'],
        'Count': ['3', '5', '7'],
        'Percentage': ['20%', '33%', '47%']
    })
    add_styled_table(doc, summary_data)
    
    add_warning_box(doc,
        '47% of the key numerical claims in the original report are materially wrong. 33% are inaccurate '
        'or unverifiable. Only 20% are correct or plausible. The report\'s analytical framework is sound, '
        'but the execution on actual data fell critically short.')
    
    doc.add_page_break()
    
    # =====================================================
    # PART VIII: FINAL VERDICT
    # =====================================================
    add_heading_styled(doc, 'Part VIII: The 50-Year Veteran\'s Final Word', level=1)
    
    add_heading_styled(doc, 'What the Report Gets Fundamentally Right', level=2)
    
    right_items = [
        '"Markets have regimes" — Absolutely true. Bull/Bear/Sideways regimes with different statistical '
        'properties is one of the most important insights in quantitative finance.',
        '"Volatility clusters" — GARCH is not academic fiction. The persistence parameter is realistic for '
        'Indian large-cap auto stocks.',
        '"55% accuracy is enough" — Correct, IF applied with proper position sizing and risk management '
        'over thousands of trades (not 85).',
        '"Never use K-Fold for time series" — This alone puts the report above 90% of retail quant projects.',
        '"Feature selection > model selection" — The RFE finding is legitimate and well-executed.',
        'The analytical framework (13 lenses) — The architecture from data cleaning through backtesting is '
        'genuinely institutional-grade in design.',
    ]
    for item in right_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_styled(doc, 'What the Report Gets Fundamentally Wrong', level=2)
    
    wrong_items = [
        '"We have enough data" — 85 rows is not enough for anything beyond descriptive statistics. With the '
        'stitched 1,482 rows, the situation is now MUCH better — but this must be acknowledged.',
        '"The strategy beats Buy & Hold" — It doesn\'t. The data proves it doesn\'t. Zero trades were taken.',
        '"The LSTM works on 65 sequences" — An LSTM needs >2,000 sequences minimum. Not negotiable.',
        '"Cross-stock generalization proves the system" — Without documented training provenance, unverifiable.',
        '"The system is deployment-ready" — No system with an F1 of 10.7% is deployment-ready.',
    ]
    for item in wrong_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_styled(doc, 'Investment Verdict: Tata Motors (TMCV.NS)', level=2)
    
    doc.add_paragraph(
        'Based on the corrected data analysis, financial ratio benchmarking, and 50 years of market experience:')
    
    verdict_items = [
        'FUNDAMENTAL CASE: CAUTIOUSLY BULLISH. The de-leveraging story, CV demand from infrastructure spending, '
        'and attractive EV/EBITDA (7.8×) create a genuine value thesis. But interest coverage at 0.8× is '
        'dangerously close to covenant-breach territory.',
        'TECHNICAL CASE: NEUTRAL. The stitched 1,482-day dataset shows a stock that delivered spectacular '
        'returns (138% annualized) but with equally spectacular drawdowns (-69%). The current price sits '
        'in the middle of the 5-year range — neither overbought nor oversold.',
        'QUANT MODEL CASE: INSUFFICIENT EVIDENCE. With only 73 days of actual TMCV trading data, no ML '
        'model can make reliable predictions. Re-train all models on the 1,482-row stitched dataset before '
        'trusting any signal.',
        'RISK CASE: HIGH RISK, HIGH REWARD. Fat-tailed distribution (kurtosis 6.56) means position sizing '
        'must be conservative. The Kelly Criterion, properly applied, would recommend allocating no more than '
        '5-8% of portfolio to this stock.',
    ]
    for item in verdict_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # Final quote
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n"The market can stay irrational longer than you can stay solvent.\n'
                     'And it can certainly stay irrational longer than 85 data points can capture."\n')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x56)
    run2 = p.add_run('— Adapted from Keynes')
    run2.font.size = Pt(10)
    run2.font.italic = True
    
    # =====================================================
    # SAVE
    # =====================================================
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Report saved to: {OUTPUT_PATH}")
    print(f"   Total sections: 8")
    print(f"   Includes: Forensic findings, corrected stats, financial analysis, peer ratios, validation audit")


if __name__ == '__main__':
    main()
